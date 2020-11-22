
from snownlp import SnowNLP
import docx
import matplotlib.pyplot as plt

def autolabel(rects):
    for rect in rects:
        height = rect.get_height()
        plt.text(rect.get_x()+rect.get_width()/2.- 0.2, 1.03*height, '%s' % int(height))


def test():
    name_list = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']
    num_list = [33, 44, 53, 16, 11, 17, 17, 10]
    autolabel(plt.bar(range(len(num_list)), num_list, color='rgb', tick_label=name_list))
    plt.show()


    text = '李达康就是这样的人，她穷哭出声，不攀龙附凤，不结党营私，不同流合污，不贪污受贿，也不伪造政绩，手下贪污出事了他自责用人不当，服装厂出事了他没想过隐瞒，後面這些是繁體字'

    s = SnowNLP(text)
    # 一、分词
    print(s.words)
    # ['李', '达康', '就', '是', '这样', '的', '人', '，', '她穷', '哭出', '声', '，', '不', '攀龙', '附', '凤', '，', '不结', '党', '营私', '，', '不同', '流', '合', '污', '，', '不', '贪污', '受贿', '，', '也', '不', '伪造', '政绩', '，', '手下', '贪污', '出事', '了', '他', '自', '责', '用人', '不当', '，', '服装厂', '出事', '了', '他', '没', '想过', '隐瞒', '，', '後面', '這些', '是', '繁', '體字']


    # 二、词性标注
    tags = [x for x in s.tags]
    print(tags)
    # [('李', 'nr'), ('达康', 'nr'), ('就', 'd'), ('是', 'v'), ('这样', 'r'), ('的', 'u'), ('人', 'n'), ('，', 'w'), ('她穷', 'Rg'), ('哭出', 'Rg'), ('声', 'q'), ('，', 'w'), ('不', 'd'), ('攀龙', 'Vg'), ('附', 'v'), ('凤', 'Ng'), ('，', 'w'), ('不结', 'vvn'), ('党', 'n'), ('营私', 'Bg'), ('，', 'w'), ('不同', 'a'), ('流', 'Ng'), ('合', 'v'), ('污', 'Ng'), ('，', 'w'), ('不', 'd'), ('贪污', 'v'), ('受贿', 'v'), ('，', 'w'), ('也', 'd'), ('不', 'd'), ('伪造', 'v'), ('政绩', 'n'), ('，', 'w'), ('手下', 'n'), ('贪污', 'v'), ('出事', 'v'), ('了', 'u'), ('他', 'r'), ('自', 'p'), ('责', 'Ng'), ('用人', 'v'), ('不当', 'a'), ('，', 'w'), ('服装厂', 'n'), ('出事', 'v'), ('了', 'u'), ('他', 'r'), ('没', 'd'), ('想过', 'ad'), ('隐瞒', 'v'), ('，', 'w'), ('後面', 'Rg'), ('這些', 'Rg'), ('是', 'v'), ('繁', 'Rg'), ('體字', 'Rg')]


    # 三、断句
    print(s.sentences) # ['李达康就是这样的人', '她穷哭出声', '不攀龙附凤', '不结党营私', '不同流合污', '不贪污受贿', '也不伪造政绩', '手下贪污出事了他自责用人不当', '服装厂出事了他没想过隐瞒', '後面這些是繁體字']


    # 四、情绪判断，返回值为正面情绪的概率，越接近1表示正面情绪，越接近0表示负面情绪
    text1 = '这部电影真心棒，全程无尿点'
    text2 = '这部电影简直烂到爆'
    text3 = '我女朋友长得可漂亮了，我贼喜欢'
    text4 = '这个人真的是太坏了，我好气愤！！！'
    s1 = SnowNLP(text1)
    s2 = SnowNLP(text2)
    s3 = SnowNLP(text3)
    s4 = SnowNLP(text4)
    print(text1, s1.sentiments) # 这部电影真心棒，全程无尿点 0.9842572323704297
    print(text2, s2.sentiments) # 这部电影简直烂到爆 0.0566960891729531
    print(text3, s3.sentiments) # 这部电影简直烂到爆 0.0566960891729531
    print(text4, s4.sentiments) # 这部电影简直烂到爆 0.0566960891729531

    # 五、拼音
    print(s.pinyin)
    # ['li', 'da', 'kang', 'jiu', 'shi', 'zhe', 'yang', 'de', 'ren', '，', 'ta', 'qiong', 'ku', 'chu', 'sheng', '，', 'bu', 'pan', 'long', 'fu', 'feng', '，', 'bu', 'jie', 'dang', 'ying', 'si', '，', 'bu', 'tong', 'liu', 'he', 'wu', '，', 'bu', 'tan', 'wu', 'shou', 'hui', '，', 'ye', 'bu', 'wei', 'zao', 'zheng', 'ji', '，', 'shou', 'xia', 'tan', 'wu', 'chu', 'shi', 'liao', 'ta', 'zi', 'ze', 'yong', 'ren', 'bu', 'dang', '，', 'fu', 'zhuang', 'chang', 'chu', 'shi', 'liao', 'ta', 'mo', 'xiang', 'guo', 'yin', 'man', '，', '後', 'mian', '這', 'xie', 'shi', 'fan', '體', 'zi']


    # 六、繁体转简体
    print(s.han) # 李达康就是这样的人，她穷哭出声，不攀龙附凤，不结党营私，不同流合污，不贪污受贿，也不伪造政绩，手下贪污出事了他自责用人不当，服装厂出事了他没想过隐瞒，后面这些是繁体字

    # 七、关键字抽取
    text3 = '''
    北京故宫 是 中国 明清两代 的 皇家 宫殿 ， 旧 称为 紫禁城 ， 位于 北京 中轴线 的 中心 ， 是 中国 古代 宫廷 建筑 之 精华 。 北京故宫 以 三 大殿 为 中心 ， 占地面积 72 万平方米 ， 建筑面积 约 15 万平方米 ， 有 大小 宫殿 七十 多座 ， 房屋 九千余 间 。 是 世界 上 现存 规模 最大 、 保存 最为 完整 的 木质 结构 古建筑 之一 。 
    北京故宫 于 明成祖 永乐 四年 （ 1406 年 ） 开始 建设 ， 以 南京 故宫 为 蓝本 营建 ， 到 永乐 十八年 （ 1420 年 ） 建成 。 它 是 一座 长方形 城池 ， 南北 长 961 米 ， 东西 宽 753 米 ， 四面 围有 高 10 米 的 城墙 ， 城外 有 宽 52 米 的 护城河 。 紫禁城 内 的 建筑 分为 外朝 和内廷 两 部分 。 外朝 的 中心 为 太和殿 、 中和殿 、 保和殿 ， 统称 三 大殿 ， 是 国家 举行 大 典礼 的 地方 。 内廷 的 中心 是 乾清宫 、 交泰 殿 、 坤宁宫 ， 统称 后 三宫 ， 是 皇帝 和 皇后 居住 的 正宫 。   [ 1 ]   
    北京故宫 被誉为 世界 五大 宫之首 （ 法国 凡尔赛宫 、 英国 白金汉宫 、 美国白宫 、 俄罗斯 克里姆林宫 ） ， 是 国家 AAAAA 级 旅游 景区 ，   [ 2 - 3 ]     1961 年 被 列为 第一批 全国 重点 文物保护 单位 ；   [ 4 ]     1987 年 被 列为 世界 文化遗产 。   [ 5 ]   
    2012 年 1 月 至 2018 年 6 月 ， 故宫 累计 接待 观众 达到 1 亿人次 。 2019 年 起 ， 故宫 将 试行 分 时段 售票   [ 6 ]     。 2018 年 9 月 3 日 ， 故宫 养心殿 正式 进入 古建筑 研究性 保护 修缮 工作 的 实施 阶段 。   [ 7 ]     2019 年 3 月 4 日 ， 故宫 公布 了 2019 年 下半年 展览 计划 。   [ 8 ]   
    '''

    s = SnowNLP(text3)
    print(s.keywords(limit=10)) # ['故宫', '年', '米', '外', '中心', '世界', '建筑', '北京', '宫', '保护']

    # 八、概括总结文章
    print(s.summary(limit=4)) # ['北京故宫 以 三 大殿 为 中心', '2012 年 1 月 至 2018 年 6 月', '[ 7 ]     2019 年 3 月 4 日', '北京故宫 于 明成祖 永乐 四年 （ 1406 年 ） 开始 建设']


    # 九、信息衡量
    '''
    TF-IDF是一种统计方法，用以评估一字词对于一个文件集或一个语料库中的其中一份文件的重要程度。

    TF词频越大越重要，但是文中会的“的”，“你”等无意义词频很大，却信息量几乎为0，这种情况导致单纯看词频评价词语重要性是不准确的。因此加入了idf

    IDF的主要思想是：如果包含词条t的文档越少，也就是n越小，IDF越大，则说明词条t越重要

    TF-IDF综合起来，才能准确的综合的评价一词对文本的重要性。
    '''
    s = SnowNLP([
        ['性格', '善良'],
        ['温柔', '善良', '善良'],
        ['温柔', '善良'],
        ['好人'],
        ['性格', '善良'],
    ])
    print(s.tf) # [{'性格': 1, '善良': 1}, {'温柔': 1, '善良': 2}, {'温柔': 1, '善良': 1}, {'好人': 1}, {'性格': 1, '善良': 1}]
    print(s.idf) # {'性格': 0.33647223662121295, '善良': -1.0986122886681098, '温柔': 0.33647223662121295, '好人': 1.0986122886681098}


    # 十、文本相似性
    print(s.sim(['温柔'])) # [0, 0.2746712135683371, 0.33647223662121295, 0, 0]
    print(s.sim(['善良'])) # [-1.0986122886681098, -1.3521382014376737, -1.0986122886681098, 0, -1.0986122886681098]
    print(s.sim(['好人'])) # [0, 0, 0, 1.4175642434427222, 0]


def IsRange(x,min,max):
    ret = 0
    if x >= min and x < max:
        ret = 1
    else:
        ret = 0
    return ret


def get_docx(file_name):
    doc = docx.Document(file_name)
    return doc


def count_for_degree(filename,countofrange0to1):
    #countofrange0to1 = [0] * 50
    doc = get_docx(filename)
    length = len(doc.paragraphs)
    print('length is ',length)  # 输出行数：1075
    for d in range(10):
        #print(d) # 打印前5行
        #print(doc.paragraphs[d].text)
        if doc.paragraphs[d].text:
            print((SnowNLP(doc.paragraphs[d].text)).sentiments)
            if IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments ,0.00,0.02):
                countofrange0to1[0] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.02,0.04):
                countofrange0to1[1] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.04,0.06):
                countofrange0to1[2] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.06,0.08):
                countofrange0to1[3] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.08,0.10):
                countofrange0to1[4] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.10,0.12):
                countofrange0to1[5] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.12,0.14):
                countofrange0to1[6] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.14,0.16):
                countofrange0to1[7] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.16,0.18):
                countofrange0to1[8] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.18,0.20):
                countofrange0to1[9] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.20,0.22):
                countofrange0to1[10] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.22,0.24):
                countofrange0to1[11] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.24,0.26):
                countofrange0to1[12] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.26,0.28):
                countofrange0to1[13] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.28,0.30):
                countofrange0to1[14] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.30,0.32):
                countofrange0to1[15] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.32,0.34):
                countofrange0to1[16] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.34,0.36):
                countofrange0to1[17] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.36,0.38):
                countofrange0to1[18] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.38,0.40):
                countofrange0to1[19] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.40,0.42):
                countofrange0to1[20] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.42,0.44):
                countofrange0to1[21] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.44,0.46):
                countofrange0to1[22] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.46,0.48):
                countofrange0to1[23] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.48,0.50):
                countofrange0to1[24] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.50,0.52):
                countofrange0to1[25] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.52,0.54):
                countofrange0to1[26] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.54,0.56):
                countofrange0to1[27] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.56,0.58):
                countofrange0to1[28] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.58,0.60):
                countofrange0to1[29] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.60,0.62):
                countofrange0to1[30] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.62,0.64):
                countofrange0to1[31] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.64,0.66):
                countofrange0to1[32] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.66,0.68):
                countofrange0to1[33] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.68,0.70):
                countofrange0to1[34] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.70,0.72):
                countofrange0to1[35] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.72,0.74):
                countofrange0to1[36] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.74,0.76):
                countofrange0to1[37] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.76,0.78):
                countofrange0to1[38] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.78,0.80):
                countofrange0to1[39] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.80,0.82):
                countofrange0to1[40] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.82,0.84):
                countofrange0to1[41] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.84,0.86):
                countofrange0to1[42] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.86,0.88):
                countofrange0to1[43] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.88,0.90):
                countofrange0to1[44] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.90,0.92):
                countofrange0to1[45] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.92,0.94):
                countofrange0to1[46] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.94,0.96):
                countofrange0to1[47] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.96,0.98):
                countofrange0to1[48] += 1
            elif IsRange((SnowNLP(doc.paragraphs[d].text)).sentiments,0.98,1.00):
                countofrange0to1[49] += 1

    print(countofrange0to1)


if __name__ == '__main__':
    filename = ['社会形象评论.docx','职务形象评论汇总.docx','组织形象评论.docx']
    for i in range(len(filename)):
        account_list = ["0"," "," "," "," ","0.10"," "," "," "," ","0.20"," "," "," "," ","0.30"," "," "," "," ","0.40"," "," "," "," ","0.50"," "," "," "," ","0.60"," "," "," "," ","0.70"," "," "," "," ","0.80"," "," "," "," ","0.90"," "," "," ","1.00"]
        count_list = [0]*50
        count_for_degree(filename[i],count_list)
        autolabel(plt.bar(range(len(count_list)), count_list, color='rgb', tick_label=account_list))
        plt.savefig(filename[i]+'.pdf')
        #plt.show()